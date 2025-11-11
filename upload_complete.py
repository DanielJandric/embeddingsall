#!/usr/bin/env python3
"""
Script complet d'upload avec Azure OCR pour documents complexes et PyPDF2 pour PDFs simples
"""

import os
import sys
import argparse
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
import hashlib
from datetime import datetime
import logging
import json
from typing import Optional, Dict, List
import time

# Configuration du logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Configuration depuis les variables d'environnement
AZURE_ENDPOINT = os.getenv("AZURE_FORM_RECOGNIZER_ENDPOINT", "https://mcpdj.cognitiveservices.azure.com/")
AZURE_KEY = os.getenv("AZURE_FORM_RECOGNIZER_KEY", "AZURE_KEY_REDACTED")
OPENAI_KEY = os.getenv("OPENAI_API_KEY", "OPENAI_KEY_REDACTED")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-3-small")
SUPABASE_URL = os.getenv("SUPABASE_URL", "https://kpfitkmaaztrjwqvockf.supabase.co")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICEROLE_KEY", "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImtwZml0a21hYXp0cmp3cXZvY2tmIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc2MjU5MDkyOCwiZXhwIjoyMDc4MTY2OTI4fQ.NYrNsMHTy-GVgyUAsiC0l1-mU-mdQUXZLs2CW-O5yAQ")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1500"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "300"))

# Installer les dépendances si nécessaire
try:
    import PyPDF2
    from supabase import create_client, Client
    from openai import OpenAI
except ImportError:
    logger.info("📦 Installing required packages...")
    os.system("python -m pip install PyPDF2 supabase openai azure-ai-formrecognizer python-docx pillow")
    import PyPDF2
    from supabase import create_client, Client
    from openai import OpenAI

class DocumentProcessor:
    """Processeur de documents avec Azure OCR et PyPDF2"""
    
    def __init__(self):
        self.supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
        self.openai_client = OpenAI(api_key=OPENAI_KEY)
        self.azure_client = None
        self._init_azure()
        
    def _init_azure(self):
        """Initialise Azure Form Recognizer si disponible"""
        try:
            from azure.ai.formrecognizer import DocumentAnalysisClient
            from azure.core.credentials import AzureKeyCredential
            
            self.azure_client = DocumentAnalysisClient(
                endpoint=AZURE_ENDPOINT,
                credential=AzureKeyCredential(AZURE_KEY)
            )
            logger.info("✅ Azure OCR initialized")
        except Exception as e:
            logger.warning(f"⚠️ Azure OCR not available: {e}")
            self.azure_client = None
    
    def extract_text_simple(self, file_path: Path) -> Optional[str]:
        """Extraction simple pour PDF avec PyPDF2"""
        try:
            if file_path.suffix.lower() == '.pdf':
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    return text if text.strip() else None
                    
            elif file_path.suffix.lower() in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
            elif file_path.suffix.lower() in ['.doc', '.docx']:
                try:
                    import docx
                    doc = docx.Document(str(file_path))
                    return '\n'.join([para.text for para in doc.paragraphs])
                except:
                    return None
            else:
                # Essayer de lire comme texte
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
        except Exception as e:
            logger.debug(f"Simple extraction failed for {file_path}: {e}")
            return None
    
    def extract_text_azure(self, file_path: Path) -> Optional[Dict]:
        """Extraction avancée avec Azure Form Recognizer"""
        if not self.azure_client:
            return None
            
        try:
            with open(file_path, 'rb') as f:
                poller = self.azure_client.begin_analyze_document(
                    "prebuilt-layout", document=f
                )
                result = poller.result()
                
            text = ""
            metadata = {
                'page_count': len(result.pages),
                'tables': [],
                'key_value_pairs': []
            }
            
            # Extraire le texte par page
            for page in result.pages:
                page_text = ""
                for line in page.lines:
                    page_text += line.content + "\n"
                text += f"\n--- Page {page.page_number} ---\n{page_text}"
            
            # Extraire les tables
            for table in result.tables:
                table_data = []
                for cell in table.cells:
                    table_data.append({
                        'row': cell.row_index,
                        'col': cell.column_index,
                        'text': cell.content
                    })
                metadata['tables'].append(table_data)
            
            # Extraire les paires clé-valeur
            if hasattr(result, 'key_value_pairs'):
                for kv in result.key_value_pairs:
                    if kv.key and kv.value:
                        metadata['key_value_pairs'].append({
                            'key': kv.key.content,
                            'value': kv.value.content
                        })
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.debug(f"Azure extraction failed for {file_path}: {e}")
            return None
    
    def generate_embedding(self, text: str) -> List[float]:
        """Génère un embedding pour un texte"""
        try:
            response = self.openai_client.embeddings.create(
                input=text[:8000],  # Limite OpenAI
                model=EMBEDDING_MODEL
            )
            return response.data[0].embedding
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            return []
    
    def chunk_text(self, text: str) -> List[str]:
        """Découpe le texte en chunks avec chevauchement"""
        if len(text) <= CHUNK_SIZE:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = min(start + CHUNK_SIZE, len(text))
            
            # Essayer de couper à un espace
            if end < len(text):
                last_space = text.rfind(' ', start, end)
                if last_space > start:
                    end = last_space
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Avancer avec chevauchement
            start = end - CHUNK_OVERLAP if end < len(text) else end
            
        return chunks
    
    def process_file(self, file_path: Path) -> Optional[Dict]:
        """Traite un fichier complet"""
        logger.info(f"📄 Processing: {file_path.name}")
        
        # Calculer le checksum
        with open(file_path, 'rb') as f:
            content = f.read()
            file_hash = hashlib.sha256(content).hexdigest()
        
        # Vérifier si déjà uploadé
        existing = self.supabase.table('documents_full').select('id').eq('checksum', file_hash).execute()
        if existing.data:
            logger.info(f"⏭️ Skipping {file_path.name} - already uploaded")
            return None
        
        # Essayer extraction simple d'abord
        text = self.extract_text_simple(file_path)
        metadata = {'extraction_method': 'simple'}
        
        # Si échec ou document complexe, utiliser Azure
        if not text or (file_path.suffix.lower() in ['.pdf', '.png', '.jpg', '.jpeg', '.tiff']):
            azure_result = self.extract_text_azure(file_path)
            if azure_result:
                text = azure_result['text']
                metadata.update(azure_result['metadata'])
                metadata['extraction_method'] = 'azure'
        
        if not text or len(text.strip()) < 10:
            logger.warning(f"⚠️ No text extracted from {file_path.name}")
            return None
        
        # Enrichir les métadonnées
        metadata.update({
            'file_name': file_path.name,
            'file_path': str(file_path),
            'file_size': len(content),
            'checksum': file_hash,
            'upload_date': datetime.now().isoformat(),
            'file_type': file_path.suffix.lower()
        })
        
        # Extraire métadonnées immobilières du nom de fichier
        name_lower = file_path.stem.lower()
        if 'bail' in name_lower or 'contrat' in name_lower:
            metadata['type_document'] = 'Contrat de bail'
            metadata['categorie'] = 'Immobilier'
        elif 'facture' in name_lower:
            metadata['type_document'] = 'Facture'
            metadata['categorie'] = 'Comptabilité'
        elif 'devis' in name_lower:
            metadata['type_document'] = 'Devis'
            metadata['categorie'] = 'Commercial'
        
        # Insérer le document principal
        doc_data = {
            'file_name': file_path.name,
            'file_path': str(file_path),
            'full_content': text,
            'file_size_bytes': len(content),
            'checksum': file_hash,
            'metadata': metadata,
            'type_document': metadata.get('type_document'),
            'categorie': metadata.get('categorie'),
            'extraction_version': '2.0',
            'last_indexed_at': datetime.now().isoformat()
        }
        
        doc_result = self.supabase.table('documents_full').insert(doc_data).execute()
        
        if not doc_result.data:
            logger.error(f"❌ Failed to insert document {file_path.name}")
            return None
        
        document_id = doc_result.data[0]['id']
        
        # Créer et insérer les chunks
        chunks = self.chunk_text(text)
        chunk_records = []
        
        for idx, chunk in enumerate(chunks):
            embedding = self.generate_embedding(chunk)
            
            chunk_data = {
                'document_id': document_id,
                'chunk_index': idx,
                'chunk_content': chunk,
                'embedding': embedding,
                'chunk_metadata': {
                    'total_chunks': len(chunks),
                    'chunk_size': len(chunk)
                }
            }
            
            # Ajouter contexte si disponible
            if idx > 0:
                chunk_data['context_before'] = chunks[idx-1][-100:]
            if idx < len(chunks) - 1:
                chunk_data['context_after'] = chunks[idx+1][:100]
            
            chunk_records.append(chunk_data)
        
        # Insérer les chunks par batch
        if chunk_records:
            self.supabase.table('document_chunks').insert(chunk_records).execute()
        
        logger.info(f"✅ Uploaded {file_path.name}: {len(chunks)} chunks")
        return {'document_id': document_id, 'chunks': len(chunks)}

def main():
    parser = argparse.ArgumentParser(description='Upload documents with Azure OCR and embeddings')
    parser.add_argument('-i', '--input', required=True, help='Input directory')
    parser.add_argument('--workers', type=int, default=3, help='Number of parallel workers')
    parser.add_argument('--dry-run', action='store_true', help='Test run without uploading')
    args = parser.parse_args()
    
    # Vérifier le répertoire d'entrée
    input_path = Path(args.input)
    if not input_path.exists():
        logger.error(f"❌ Directory not found: {input_path}")
        sys.exit(1)
    
    # Trouver tous les fichiers
    extensions = ['.pdf', '.txt', '.md', '.doc', '.docx', '.png', '.jpg', '.jpeg', '.tiff']
    files = []
    for ext in extensions:
        files.extend(input_path.rglob(f'*{ext}'))
    
    logger.info(f"📁 Found {len(files)} files to process")
    
    if not files:
        logger.warning("No supported files found!")
        return
    
    if args.dry_run:
        logger.info("🔍 DRY RUN - Files that would be processed:")
        for f in files[:20]:  # Afficher les 20 premiers
            logger.info(f"  - {f.name}")
        if len(files) > 20:
            logger.info(f"  ... and {len(files) - 20} more files")
        return
    
    # Initialiser le processeur
    processor = DocumentProcessor()
    
    # Traiter les fichiers en parallèle
    success_count = 0
    error_count = 0
    total_chunks = 0
    
    with ThreadPoolExecutor(max_workers=args.workers) as executor:
        futures = {executor.submit(processor.process_file, f): f for f in files}
        
        for future in as_completed(futures):
            try:
                result = future.result()
                if result:
                    success_count += 1
                    total_chunks += result['chunks']
                else:
                    error_count += 1
            except Exception as e:
                logger.error(f"❌ Processing error: {e}")
                error_count += 1
    
    # Résumé
    logger.info(f"\n{'='*60}")
    logger.info(f"📊 UPLOAD COMPLETE")
    logger.info(f"✅ Successfully uploaded: {success_count} files")
    logger.info(f"📝 Total chunks created: {total_chunks}")
    logger.info(f"❌ Failed/Skipped: {error_count} files")
    logger.info(f"{'='*60}")

if __name__ == "__main__":
    main()
