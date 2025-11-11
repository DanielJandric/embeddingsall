#!/usr/bin/env python3
"""
Upload avec Azure OCR obligatoire pour extraction maximale
"""
import os
import sys
from pathlib import Path
import hashlib
from datetime import datetime
import time

# Exige les variables d'environnement (ne rien mettre en dur dans le code)
required_env = [
    "OPENAI_API_KEY",
    "SUPABASE_URL",
]
missing = [k for k in required_env if not os.getenv(k)]
if missing:
    print(f"❌ Missing environment variables: {', '.join(missing)}")
    print("Set them in PowerShell before running, e.g.:")
    print('$env:OPENAI_API_KEY="sk-..."')
    print('$env:SUPABASE_URL="https://....supabase.co"')
    sys.exit(1)

# Chunks longs avec beaucoup de contexte
CHUNK_SIZE = 2500
CHUNK_OVERLAP = 500

print("📦 Installing packages...")
os.system("pip install -q supabase openai azure-ai-formrecognizer azure-core openpyxl pandas")

from supabase import create_client
from openai import OpenAI
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential

# Clients
supabase = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_KEY"])
openai_client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

# Azure OCR client
azure_client = DocumentAnalysisClient(
    endpoint=os.environ["AZURE_FORM_RECOGNIZER_ENDPOINT"],
    credential=AzureKeyCredential(os.environ["AZURE_FORM_RECOGNIZER_KEY"])
)

def extract_text_azure(file_path):
    """Utilise Azure OCR pour extraire TOUT le texte"""
    print(f"  🔍 Azure OCR processing...")
    try:
        with open(file_path, 'rb') as f:
            # Utiliser prebuilt-read pour OCR maximal
            poller = azure_client.begin_analyze_document(
                "prebuilt-read",  # Modèle optimisé pour OCR
                document=f
            )
            result = poller.result()
        
        # Extraire tout le texte
        full_text = ""
        
        # Méthode 1: Utiliser le contenu direct si disponible
        if hasattr(result, 'content') and result.content:
            full_text = result.content
            print(f"  ✅ Extracted via content: {len(full_text)} chars")
        
        # Méthode 2: Parcourir les pages
        if not full_text and hasattr(result, 'pages'):
            for page in result.pages:
                page_text = ""
                
                # Extraire les lignes
                if hasattr(page, 'lines'):
                    for line in page.lines:
                        if hasattr(line, 'content'):
                            page_text += line.content + "\n"
                
                # Extraire les mots si pas de lignes
                if not page_text and hasattr(page, 'words'):
                    for word in page.words:
                        if hasattr(word, 'content'):
                            page_text += word.content + " "
                    page_text += "\n"
                
                if page_text:
                    full_text += f"\n--- Page {page.page_number} ---\n{page_text}"
        
        # Méthode 3: Utiliser les paragraphes si disponibles
        if not full_text and hasattr(result, 'paragraphs'):
            for para in result.paragraphs:
                if hasattr(para, 'content'):
                    full_text += para.content + "\n\n"
        
        # Extraire les tables aussi
        table_text = ""
        if hasattr(result, 'tables'):
            for table_idx, table in enumerate(result.tables):
                table_text += f"\n[Table {table_idx + 1}]\n"
                current_row = -1
                row_text = ""
                
                for cell in sorted(table.cells, key=lambda x: (x.row_index, x.column_index)):
                    if cell.row_index != current_row:
                        if row_text:
                            table_text += row_text + "\n"
                        row_text = ""
                        current_row = cell.row_index
                    row_text += f"{cell.content} | "
                
                if row_text:
                    table_text += row_text + "\n"
        
        if table_text:
            full_text += "\n" + table_text
        
        if full_text:
            print(f"  ✅ Azure OCR: {len(full_text)} chars extracted")
            return full_text
        else:
            print(f"  ⚠️ Azure OCR: No text found")
            return None
            
    except Exception as e:
        print(f"  ❌ Azure OCR error: {e}")
        # Fallback to simple extraction
        return extract_text_simple(file_path)

def extract_text_simple(file_path):
    """Extraction simple en fallback"""
    try:
        if file_path.suffix.lower() in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif file_path.suffix.lower() in ['.doc', '.docx']:
            try:
                import docx
                doc = docx.Document(str(file_path))
                return '\n'.join([p.text for p in doc.paragraphs])
            except:
                pass
        elif file_path.suffix.lower() in ['.xlsx', '.xls']:
            # Extraction Excel
            try:
                import pandas as pd
                import openpyxl
                
                # Lire toutes les feuilles
                excel_file = pd.ExcelFile(file_path)
                full_text = f"=== EXCEL FILE: {file_path.name} ===\n\n"
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    full_text += f"\n--- SHEET: {sheet_name} ---\n"
                    full_text += f"Dimensions: {df.shape[0]} rows x {df.shape[1]} columns\n\n"
                    
                    # Convertir en texte structuré
                    # En-têtes
                    full_text += "COLUMNS: " + " | ".join(str(col) for col in df.columns) + "\n\n"
                    
                    # Données (limiter aux 1000 premières lignes pour éviter trop de volume)
                    for idx, row in df.head(1000).iterrows():
                        row_text = f"Row {idx+1}: "
                        row_values = []
                        for col in df.columns:
                            val = row[col]
                            if pd.notna(val):
                                row_values.append(f"{col}={val}")
                        row_text += " | ".join(row_values)
                        full_text += row_text + "\n"
                    
                    # Résumé statistique pour colonnes numériques
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        full_text += "\n--- STATISTICS ---\n"
                        for col in numeric_cols:
                            if df[col].notna().any():
                                full_text += f"{col}: min={df[col].min():.2f}, max={df[col].max():.2f}, mean={df[col].mean():.2f}\n"
                
                return full_text
            except Exception as e:
                print(f"  ⚠️ Excel extraction error: {e}")
                return None
        elif file_path.suffix.lower() == '.csv':
            # CSV simple
            try:
                import pandas as pd
                df = pd.read_csv(file_path, encoding='utf-8', on_bad_lines='skip')
                text = f"=== CSV FILE: {file_path.name} ===\n"
                text += f"Shape: {df.shape[0]} rows x {df.shape[1]} columns\n\n"
                text += df.to_string(max_rows=1000)
                return text
            except:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
    except:
        pass
    return None

def chunk_text(text):
    """Découpe en chunks longs avec overlap"""
    if len(text) <= CHUNK_SIZE:
        return [text]
    
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))
        chunk = text[start:end]
        chunks.append(chunk)
        start = end - CHUNK_OVERLAP if end < len(text) else end
    return chunks

def process_file(file_path):
    """Traite un fichier avec Azure OCR"""
    print(f"\n📄 {file_path.name}")
    
    # Checksum
    with open(file_path, 'rb') as f:
        content = f.read()
        checksum = hashlib.sha256(content).hexdigest()
    
    # Vérifier si déjà uploadé
    existing = supabase.table('documents_full').select('id').eq('checksum', checksum).execute()
    if existing.data:
        print(f"  ⏭️ Already uploaded")
        return
    
    # Stratégie d'extraction selon le type
    text = None
    if file_path.suffix.lower() in ['.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
        # Azure OCR pour PDFs et images
        text = extract_text_azure(file_path)
    elif file_path.suffix.lower() in ['.xlsx', '.xls', '.csv']:
        # Extraction directe pour Excel/CSV
        text = extract_text_simple(file_path)
    else:
        # Autres formats : essayer simple d'abord, puis Azure
        text = extract_text_simple(file_path)
        if not text:
            text = extract_text_azure(file_path)
    
    if not text or len(text.strip()) < 10:
        print(f"  ⚠️ No text extracted - skipping")
        return
    
    print(f"  📝 Text extracted: {len(text)} chars")
    
    # Document principal
    doc_data = {
        'file_name': file_path.name,
        'file_path': str(file_path),
        'full_content': text,
        'file_size_bytes': len(content),
        'checksum': checksum,
        'metadata': {
            'upload_date': datetime.now().isoformat(),
            'extraction_method': 'azure_ocr',
            'chunk_config': f'{CHUNK_SIZE}/{CHUNK_OVERLAP}'
        }
    }
    
    # Ajouter type de document si détecté
    name_lower = file_path.stem.lower()
    if 'bail' in name_lower or 'contrat' in name_lower:
        doc_data['type_document'] = 'Contrat'
        doc_data['categorie'] = 'Immobilier'
    elif 'facture' in name_lower:
        doc_data['type_document'] = 'Facture'
        doc_data['categorie'] = 'Comptabilité'
    
    doc_result = supabase.table('documents_full').insert(doc_data).execute()
    if not doc_result.data:
        print(f"  ❌ Insert failed")
        return
    
    doc_id = doc_result.data[0]['id']
    
    # Chunks
    chunks = chunk_text(text)
    print(f"  📊 Creating {len(chunks)} chunks...")
    
    for idx, chunk in enumerate(chunks):
        # Embedding
        try:
            response = openai_client.embeddings.create(
                input=chunk[:8000],
                model="text-embedding-3-small"
            )
            embedding = response.data[0].embedding
        except Exception as e:
            print(f"  ⚠️ Embedding error: {e}")
            embedding = []
        
        # Contexte enrichi
        chunk_data = {
            'document_id': doc_id,
            'chunk_index': idx,
            'chunk_content': chunk,
            'chunk_size': len(chunk),  # Colonne directe pour Supabase
            'embedding': embedding,
            'chunk_metadata': {
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
        }
        
        # Ajouter contexte avant/après
        if idx > 0:
            chunk_data['context_before'] = chunks[idx-1][-300:]  # 300 chars de contexte
        if idx < len(chunks)-1:
            chunk_data['context_after'] = chunks[idx+1][:300]
        
        supabase.table('document_chunks').insert(chunk_data).execute()
        
        # Petit délai pour ne pas surcharger
        if idx % 10 == 0:
            time.sleep(0.1)
    
    print(f"  ✅ Success: {len(chunks)} chunks uploaded")

from concurrent.futures import ThreadPoolExecutor, as_completed
import threading

# Thread-safe counters
success_lock = threading.Lock()
failed_lock = threading.Lock()
success_count = 0
failed_count = 0

def process_file_wrapper(file_path):
    """Wrapper pour gérer les compteurs thread-safe"""
    global success_count, failed_count
    try:
        process_file(file_path)
        with success_lock:
            success_count += 1
        return True
    except Exception as e:
        print(f"❌ Error processing {file_path.name}: {e}")
        with failed_lock:
            failed_count += 1
        return False

# Main
input_dir = Path(r"C:\OneDriveExport")

# Tous les formats supportés
extensions = ['.pdf', '.txt', '.doc', '.docx', '.md', 
              '.png', '.jpg', '.jpeg', '.tiff', '.bmp',
              '.xlsx', '.xls', '.csv']  # Excel et CSV ajoutés

files = []
for ext in extensions:
    files.extend(input_dir.rglob(f'*{ext}'))

print(f"\n🚀 AZURE OCR UPLOAD WITH 3 WORKERS")
print(f"📁 Found {len(files)} files in {input_dir}")
print(f"🔧 Config: Chunks={CHUNK_SIZE}, Overlap={CHUNK_OVERLAP}")
print(f"⚡ Workers: 3 parallel threads")
print(f"{'='*60}\n")

# Traiter avec 3 workers en parallèle
with ThreadPoolExecutor(max_workers=3) as executor:
    futures = {executor.submit(process_file_wrapper, f): f for f in files}
    
    for future in as_completed(futures):
        file = futures[future]
        try:
            result = future.result()
        except Exception as e:
            print(f"❌ Thread error for {file.name}: {e}")

print(f"\n{'='*60}")
print(f"✅ SUCCESS: {success_count} files")
print(f"❌ FAILED: {failed_count} files")
print(f"{'='*60}")
